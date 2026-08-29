"""Extract plain text from an uploaded resume file so it can be chunked and
embedded. Supports PDF, DOCX, and plain text — the formats resumes actually
come in.
"""

from io import BytesIO


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    elif ext == "docx":
        return _extract_docx(file_bytes)
    elif ext in ("txt", "md"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported resume file type '.{ext}'. Please upload a PDF, DOCX, or TXT file."
        )


def _extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if not text:
        raise ValueError(
            "Could not extract any text from this PDF — it may be a scanned "
            "image rather than a text-based PDF."
        )
    return text


def _extract_docx(file_bytes: bytes) -> str:
    import docx

    document = docx.Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Also pull text out of any tables (some resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("Could not extract any text from this DOCX file.")
    return text
